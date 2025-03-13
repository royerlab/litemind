import os
import mimetypes
from typing import List, Optional, Union

from arbol import aprint

from litemind.agent.messages.message import Message
from litemind.agent.messages.message_block_type import BlockType
from litemind.agent.augmentations.augmentation_base import Document
from litemind.agent.augmentations.hierarchical_document import HierarchicalDocument, create_document_hierarchy
from litemind.utils.normalise_uri_to_local_file_path import uri_to_local_file_path


class DocumentLoader:
    """
    Utility class for loading documents from various sources.
    """
    
    @staticmethod
    def load_from_file(
        file_path: str, 
        split_documents: bool = True, 
        chunk_size: int = 1000, 
        chunk_overlap: int = 200
    ) -> Union[HierarchicalDocument, List[Document]]:
        """
        Load a document from a file.
        
        Parameters
        ----------
        file_path: str
            The path to the file.
        split_documents: bool
            Whether to split the document into chunks.
        chunk_size: int
            The size of each chunk.
        chunk_overlap: int
            The overlap between consecutive chunks.
            
        Returns
        -------
        Union[HierarchicalDocument, List[Document]]
            Either a hierarchical document or a list of documents.
        """
        # Determine file type
        mime_type, _ = mimetypes.guess_type(file_path)
        
        if mime_type is None:
            mime_type = "text/plain"
        
        # Read the file content
        text = ""
        metadata = {"source": file_path}
        
        aprint(f"Loading file: {file_path} (mime type: {mime_type})")
        
        if mime_type.startswith("text/"):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        elif mime_type == "application/pdf":
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                text = ""
                for page_num, page in enumerate(doc):
                    text += f"Page {page_num + 1}:\n{page.get_text()}\n\n"
            except ImportError:
                raise ImportError("PyMuPDF not installed. Please install it with 'pip install pymupdf'.")
        elif mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
            try:
                import docx
                doc = docx.Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
            except ImportError:
                raise ImportError("python-docx not installed. Please install it with 'pip install python-docx'.")
        else:
            raise ValueError(f"Unsupported file type: {mime_type}")
        
        # Process the text
        if split_documents:
            hierarchy = create_document_hierarchy(
                text, 
                max_chunk_size=chunk_size, 
                chunk_overlap=chunk_overlap
            )
            # Add metadata to the root document
            hierarchy.metadata.update(metadata)
            return hierarchy
        else:
            return [Document(content=text, metadata=metadata)]
    
    @staticmethod
    def load_from_directory(
        directory_path: str,
        recursive: bool = True,
        split_documents: bool = True,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        include_extensions: Optional[List[str]] = None,
    ) -> List[Union[HierarchicalDocument, Document]]:
        """
        Load documents from a directory.
        
        Parameters
        ----------
        directory_path: str
            The path to the directory.
        recursive: bool
            Whether to recursively process subdirectories.
        split_documents: bool
            Whether to split documents into chunks.
        chunk_size: int
            The size of each chunk.
        chunk_overlap: int
            The overlap between consecutive chunks.
        include_extensions: Optional[List[str]]
            List of file extensions to include (e.g., [".txt", ".pdf"]).
            If None, all supported files will be processed.
            
        Returns
        -------
        List[Union[HierarchicalDocument, Document]]
            A list of documents.
        """
        documents = []
        
        for root, dirs, files in os.walk(directory_path):
            for file in files:
                file_path = os.path.join(root, file)
                
                # Check extension
                _, ext = os.path.splitext(file)
                if include_extensions and ext.lower() not in include_extensions:
                    continue
                
                try:
                    doc = DocumentLoader.load_from_file(
                        file_path,
                        split_documents=split_documents,
                        chunk_size=chunk_size,
                        chunk_overlap=chunk_overlap,
                    )
                    
                    if isinstance(doc, list):
                        documents.extend(doc)
                    else:
                        documents.append(doc)
                except Exception as e:
                    aprint(f"Error loading {file_path}: {e}")
            
            # If not recursive, break after processing the top directory
            if not recursive:
                break
        
        return documents
    
    @staticmethod
    def load_from_message(
        message: Message,
        extract_text: bool = True,
        split_documents: bool = True,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ) -> List[Union[HierarchicalDocument, Document]]:
        """
        Load documents from a litemind Message object.
        
        Parameters
        ----------
        message: Message
            The message to process.
        extract_text: bool
            Whether to extract text from non-text blocks.
        split_documents: bool
            Whether to split documents into chunks.
        chunk_size: int
            The size of each chunk.
        chunk_overlap: int
            The overlap between consecutive chunks.
            
        Returns
        -------
        List[Union[HierarchicalDocument, Document]]
            A list of documents.
        """
        documents = []
        
        # Process text blocks
        text_content = message.to_plain_text()
        if text_content:
            if split_documents:
                hierarchy = create_document_hierarchy(
                    text_content,
                    max_chunk_size=chunk_size,
                    chunk_overlap=chunk_overlap,
                )
                documents.append(hierarchy)
            else:
                documents.append(Document(content=text_content))
        
        # Process document blocks if needed
        if extract_text:
            for block in message.blocks:
                if block.block_type == BlockType.Document:
                    try:
                        doc = DocumentLoader.load_from_file(
                            uri_to_local_file_path(block.content),
                            split_documents=split_documents,
                            chunk_size=chunk_size,
                            chunk_overlap=chunk_overlap,
                        )
                        
                        if isinstance(doc, list):
                            documents.extend(doc)
                        else:
                            documents.append(doc)
                    except Exception as e:
                        aprint(f"Error loading document block: {e}")
        
        return documents
